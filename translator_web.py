import streamlit as st
from google.cloud import translate
import os, io, csv, time, tempfile, re
from typing import List, Optional
from docx import Document
import pandas as pd
import speech_recognition as sr
import pyttsx3
import logging

# ---------------------------
# GOOGLE CLOUD AUTH
# ---------------------------
PROJECT_ID = "translator-demo-473714"
LOCATION = "global"
client = translate.TranslationServiceClient()
parent = f"projects/{PROJECT_ID}/locations/{LOCATION}"

# ---------------------------
# SPEECH ENGINE SETUP
# ---------------------------
engine = pyttsx3.init()
voices = engine.getProperty("voices")
engine.setProperty("voice", voices[0].id)  # default male
engine.setProperty("rate", 150)

# ---------------------------
# LOGGING
# ---------------------------
logging.basicConfig(filename='translator_errors.log', level=logging.ERROR)


# ---------------------------
# TRANSLATION FUNCTION
# ---------------------------
def translate_text(texts: List[str], target_language: str, source_language: Optional[str] = None,
                   mime_type: str = "text/plain") -> List[dict]:
    try:
        request = {
            "parent": parent,
            "contents": texts,
            "mime_type": mime_type,
            "target_language_code": target_language,
        }
        if source_language:
            request["source_language_code"] = source_language
        response = client.translate_text(**request)
        results = []
        for translation in response.translations:
            results.append({
                "translatedText": translation.translated_text,
                "detectedLanguageCode": getattr(translation, "detected_language_code", None),
            })
        return results
    except Exception as e:
        logging.error(f"{time.ctime()}: {str(e)}")
        st.error(f"Translation API error: {e}")
        return [{"translatedText": "", "error": str(e)} for _ in texts]


# ---------------------------
# FILE HANDLING
# ---------------------------
def read_file(uploaded_file):
    if uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8").splitlines()
    elif uploaded_file.name.endswith(".csv"):
        return [row[0] for row in csv.reader(io.StringIO(uploaded_file.getvalue().decode("utf-8")))]
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return [para.text for para in doc.paragraphs if para.text.strip()]
    else:
        st.error("Unsupported file type. Supported: TXT, CSV, DOCX")
        return []


def save_to_docx(lines, translations, filename="translated.docx"):
    doc = Document()
    for inp, trans in zip(lines, translations):
        doc.add_paragraph(f"Original: {inp}")
        doc.add_paragraph(f"Translated: {trans['translatedText']}")
        doc.add_paragraph("---")
    doc.save(filename)
    return filename


# ✅ FIXED: save speech as .wav in temp file
def save_speech(text):
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
    engine.save_to_file(text, tmp_file.name)
    engine.runAndWait()
    return tmp_file.name


# ---------------------------
# LANGUAGES DICTIONARY
# ---------------------------
LANGUAGES = {
    "Afrikaans": "af", "Albanian": "sq", "Amharic": "am", "Arabic": "ar", "Armenian": "hy",
    "Azerbaijani": "az", "Basque": "eu", "Belarusian": "be", "Bengali": "bn", "Bosnian": "bs",
    "Bulgarian": "bg", "Catalan": "ca", "Cebuano": "ceb", "Chinese (Simplified)": "zh-CN",
    "Chinese (Traditional)": "zh-TW", "Corsican": "co", "Croatian": "hr", "Czech": "cs",
    "Danish": "da", "Dutch": "nl", "English": "en", "Esperanto": "eo", "Estonian": "et",
    "Finnish": "fi", "French": "fr", "Frisian": "fy", "Galician": "gl", "Georgian": "ka",
    "German": "de", "Greek": "el", "Gujarati": "gu", "Haitian Creole": "ht", "Hausa": "ha",
    "Hawaiian": "haw", "Hebrew": "he", "Hindi": "hi", "Hmong": "hmn", "Hungarian": "hu",
    "Icelandic": "is", "Igbo": "ig", "Indonesian": "id", "Irish": "ga", "Italian": "it",
    "Japanese": "ja", "Javanese": "jv", "Kannada": "kn", "Kazakh": "kk", "Khmer": "km",
    "Kinyarwanda": "rw", "Korean": "ko", "Kurdish": "ku", "Kyrgyz": "ky", "Lao": "lo",
    "Latin": "la", "Latvian": "lv", "Lithuanian": "lt", "Luxembourgish": "lb", "Macedonian": "mk",
    "Malay": "ms", "Malayalam": "ml", "Maltese": "mt", "Maori": "mi", "Marathi": "mr",
    "Mongolian": "mn", "Myanmar (Burmese)": "my", "Nepali": "ne", "Norwegian": "no",
    "Nyanja (Chichewa)": "ny", "Odia (Oriya)": "or", "Pashto": "ps", "Persian": "fa",
    "Polish": "pl", "Portuguese": "pt", "Punjabi": "pa", "Romanian": "ro", "Russian": "ru",
    "Samoan": "sm", "Scots Gaelic": "gd", "Serbian": "sr", "Sesotho": "st", "Shona": "sn",
    "Sindhi": "sd", "Sinhala": "si", "Slovak": "sk", "Slovenian": "sl", "Somali": "so",
    "Spanish": "es", "Sundanese": "su", "Swahili": "sw", "Swedish": "sv", "Tagalog (Filipino)": "tl",
    "Tajik": "tg", "Tamil": "ta", "Tatar": "tt", "Telugu": "te", "Thai": "th", "Turkish": "tr",
    "Turkmen": "tk", "Ukrainian": "uk", "Urdu": "ur", "Uyghur": "ug", "Uzbek": "uz",
    "Vietnamese": "vi", "Welsh": "cy", "Xhosa": "xh", "Yiddish": "yi", "Yoruba": "yo", "Zulu": "zu"
}


# ---------------------------
# STREAMLIT UI SETUP
# ---------------------------
st.set_page_config(page_title="🌍 Advanced Translator App", layout="wide")
st.title("🌍 AI-Powered Translator with Voice & File Support")

# Sidebar Settings
st.sidebar.header("⚙️ Settings")
source_lang_name = st.sidebar.selectbox("Source Language", ["Auto Detect"] + list(LANGUAGES.keys()))
target_lang_name = st.sidebar.selectbox("Target Language", list(LANGUAGES.keys()), index=1)
source_lang = None if source_lang_name == "Auto Detect" else LANGUAGES[source_lang_name]
target_lang = LANGUAGES[target_lang_name]

voice_choice = st.sidebar.radio("🔊 Voice", ["Male", "Female"], horizontal=True)
rate = st.sidebar.slider("Speech Speed", 100, 250, 150)
if voice_choice == "Female" and len(voices) > 1:
    engine.setProperty("voice", voices[1].id)
else:
    engine.setProperty("voice", voices[0].id)
engine.setProperty("rate", rate)

# Session history
if "history" not in st.session_state:
    st.session_state.history = []

# Tabs
tab1, tab2, tab3 = st.tabs(["✍️ Text Input", "📂 File Upload", "🎙️ Conversation Mode"])

# --- Text Input ---
with tab1:
    user_text = st.text_area("Enter text to translate", height=150)
    if st.button("Translate Text"):
        if user_text.strip():
            sentences = re.split(r'(?<=[.!?]) +', user_text)
            results = translate_text(sentences, target_lang, source_lang)
            translated_text = " ".join([r['translatedText'] for r in results])

            st.success(f"✅ Translated: {translated_text}")
            detected_langs = set([r['detectedLanguageCode'] for r in results if r.get('detectedLanguageCode')])
            st.info(f"Detected Language(s): {', '.join(detected_langs)}")

            st.session_state.history.append({"input": user_text, "translatedText": translated_text,
                                             "detectedLanguageCode": ", ".join(detected_langs)})

            if st.button("🔊 Speak Result", key=f"tts_{time.time()}"):
                audio_file = save_speech(translated_text)
                st.audio(audio_file, format="audio/wav")

# --- File Upload ---
with tab2:
    uploaded_files = st.file_uploader("Drag & Drop your file(s)", type=["txt", "csv", "docx"],
                                      accept_multiple_files=True)
    if uploaded_files:
        for uploaded_file in uploaded_files:
            lines = read_file(uploaded_file)
            if st.button(f"Translate File: {uploaded_file.name}"):
                results = translate_text(lines, target_lang, source_lang)
                for inp, r in zip(lines, results):
                    st.write(f"**Input:** {inp}")
                    st.success(f"**Translated:** {r['translatedText']}")
                    st.caption(f"Detected: {r.get('detectedLanguageCode')}")
                    st.session_state.history.append({"input": inp, **r})

                # Download Options
                docx_file = save_to_docx(lines, results)
                st.download_button("⬇️ Download Translated DOCX", data=open(docx_file, "rb").read(),
                                   file_name="translated.docx")
                csv_file = pd.DataFrame({"Original": lines, "Translated": [r['translatedText'] for r in results]})
                st.download_button("⬇️ Download Translated CSV", data=csv_file.to_csv(index=False).encode("utf-8"),
                                   file_name="translated.csv")

# --- Live Conversation ---
with tab3:
    st.write("🎤 Speak and get instant translation")
    if st.button("🎙️ Start Talking"):
        recognizer = sr.Recognizer()
        with sr.Microphone() as mic:
            st.info("Listening... Speak now!")
            try:
                audio = recognizer.listen(mic, timeout=5, phrase_time_limit=7)
                spoken_text = recognizer.recognize_google(audio, language=source_lang or "en")
                st.success(f"✅ You said: {spoken_text}")

                result = translate_text([spoken_text], target_lang, source_lang)[0]
                translated = result["translatedText"]
                st.success(f"🌍 Translation: {translated}")
                st.caption(f"Detected: {result.get('detectedLanguageCode')}")

                audio_file = save_speech(translated)
                st.audio(audio_file, format="audio/wav")

                st.session_state.history.append({"input": spoken_text, **result})

            except sr.UnknownValueError:
                st.error("❌ Could not understand speech")
            except sr.RequestError:
                st.error("❌ Speech recognition service unavailable")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                logging.error(str(e))

# --- Translation History ---
if st.session_state.history:
    st.markdown("---")
    st.subheader("🕒 Translation History")
    df = pd.DataFrame(st.session_state.history)

    search_text = st.text_input("Search history")
    if search_text:
        df = df[df.apply(lambda row: row.astype(str).str.contains(search_text, case=False).any(), axis=1)]

    st.dataframe(df)

    st.download_button(
        label="⬇️ Download History CSV",
        data=df.to_csv(index=False).encode("utf-8"),
        file_name="translation_history.csv",
        mime="text/csv"
    )

    if st.button("🗑️ Clear History"):
        st.session_state.history.clear()
        st.experimental_rerun()
